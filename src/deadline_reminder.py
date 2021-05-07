import datetime as dt
import docx
from docx.opc.exceptions import PackageNotFoundError
import logging
from dotenv import dotenv_values
import os.path
import pandas as pd


def print_deadline_reminder():
    if os.path.isfile("historical_data/word_count.csv"):
        df = pd.read_csv("historical_data/word_count.csv")
        df = df.set_index("date")
    else:
        df = pd.DataFrame(columns=["date", "words", "target"])
        df = df.set_index("date")


    report_path = dotenv_values()['REPORT_PATH']

    project_start_date = dt.datetime(2021, 2, 2)
    report_start_date = dt.datetime(2021, 4, 1)
    current_date = dt.datetime.today()
    if current_date.hour >= 2:
        end_of_current_date = dt.datetime.combine(dt.date.today(), dt.datetime.min.time()) + dt.timedelta(days=1)
    else:
        end_of_current_date = dt.datetime.combine(dt.date.today(), dt.datetime.min.time())

    deadline = dt.datetime(2021, 5, 7)
    word_limit = 10500
    try:
        report = docx.Document(report_path)
        full_text = []
        first_chap_found = False
        for paragraph in report.paragraphs:
            if "INTRODUCTION" in paragraph.text:
                first_chap_found = True
            if "REFERENCES" in paragraph.text:
                break
            if first_chap_found:
                full_text.append(paragraph.text)

        for table in report.tables:
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    text = cell.text
                    full_text.append(text)

        string = '\n'.join(full_text)
        string = string.replace('\n', ' ')
        word_count = len(list(filter(None, string.split(' ')))) - 754

        total_difference = deadline - project_start_date
        report_difference = deadline - report_start_date
        current_report_progress = current_date - report_start_date
        current_project_progress = current_date - project_start_date
        report_time_progress_pct = round((current_report_progress / report_difference) * 100, 2)
        project_time_progress_pct = round((current_project_progress / total_difference) * 100, 2)
        target_word_count = round(word_limit * ((end_of_current_date - report_start_date) / report_difference))
        logging.info("*---------------------- REPORT STATUS ----------------------*")
        logging.info(f"{(deadline - current_date).days} days left until deadline. ({project_time_progress_pct}%)")
        logging.info(f"Time progress since started report: {report_time_progress_pct}%")
        logging.info(f"Target word count: {target_word_count} ({round(target_word_count / word_limit * 100, 2)}%)")
        logging.info(f"Current word count: {word_count} ({round(word_count / word_limit * 100, 2)}%)")
        df.loc[dt.datetime.strftime(current_date, "%d/%m/%Y")] = (word_count, target_word_count)
        df.to_csv("historical_data/word_count.csv")
        if word_count >= target_word_count:
            logging.info(f"YOU ARE {word_count - target_word_count} WORDS ABOVE TARGET")
            logging.info("*-----------------------------------------------------------*\n\n")
        else:
            logging.warning(f"YOU ARE {target_word_count - word_count} WORDS BELOW TARGET")
            logging.warning(f"CANNOT RUN UNTIL ABOVE THE TARGET FOR TODAY")
            logging.info("*-----------------------------------------------------------*\n\n")
            exit()
    except PackageNotFoundError as err:
        logging.warning("*----- Word document is open, cannot get wordcount pct -----*\n\n")


